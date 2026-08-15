import io
import json
import os
import re
import docx
import difflib
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Cm, Pt
from pypdf import PdfReader
import streamlit as st

# Thử import các SDK của Google Gemini
try:
    from google import genai
except ImportError:
    genai = None

try:
    import google.generativeai as genai_legacy
except ImportError:
    genai_legacy = None

# ---------------------------------------------------------
# 1. CẤU HÌNH TRANG STREAMLIT & TIÊU ĐỀ
# ---------------------------------------------------------
st.set_page_config(
    page_title="Đại học Y Hà Nội - Hệ thống Kiểm tra định dạng luận văn",
    page_icon="🎓",
    layout="wide"
)

col_logo, col_title = st.columns([1, 5])

with col_logo:
    if os.path.exists("logo_hmu.png"):
        st.image("logo_hmu.png", width=125)
    else:
        st.markdown("<h1 style='font-size: 80px; margin: 0;'>🎓</h1>", unsafe_allow_html=True)

with col_title:
    st.markdown(
        """
        <div style="line-height: 1.2; margin-top: 5px;">
            <p style="font-size: 16px; font-weight: 600; color: #666666; margin: 0; text-transform: uppercase; letter-spacing: 0.5px;">
                Trường Đại học Y Hà Nội
            </p>
            <h1 style="font-size: 28px; font-weight: 700; color: #ad171c; margin: 2px 0 8px 0; padding: 0;">
                Trung tâm Khảo thí & ĐBCLGD — Bộ môn Mắt - Khúc xạ nhãn khoa
            </h1>
            <h3 style="font-size: 18px; font-weight: 500; color: #333333; margin: 0; padding: 0;">
                🔬 Hệ thống Kiểm tra và Sửa định dạng Luận văn tự động
            </h3>
        </div>
        """,
        unsafe_allow_html=True
    )

st.caption(
    "Hệ thống tự động kiểm tra quy định, bảo toàn trang bìa/khung viền/logo, xử lý mục đạo đức nghiên cứu và chuẩn hóa file Word theo đúng quy chế."
)

API_KEY = st.secrets.get("GEMINI_API_KEY", os.environ.get("GEMINI_API_KEY", ""))

# ---------------------------------------------------------
# 2. HÀM ĐỌC TOÀN BỘ VĂN BẢN TỪ FILE PDF QUY ĐỊNH
# ---------------------------------------------------------
def extract_raw_text_from_pdf(pdf_path):
    if not os.path.exists(pdf_path):
        return ""
    try:
        reader = PdfReader(pdf_path)
        full_text = []
        for idx, page in enumerate(reader.pages):
            txt = page.extract_text()
            if txt:
                full_text.append(f"--- TRANG {idx+1} ---\n{txt}")
        return "\n".join(full_text).strip()
    except Exception as e:
        st.error(f"Lỗi đọc file PDF: {e}")
        return ""

# ---------------------------------------------------------
# 3. HÀM AI GEMINI PHÂN TÍCH PDF
# ---------------------------------------------------------
def analyze_rules_with_gemini(pdf_text, api_key):
    default_rules = {
        "font_name": "Times New Roman",
        "font_size": 13.0,
        "line_spacing": 1.5,
        "margin_top": 3.5,
        "margin_bottom": 3.0,
        "margin_left": 3.0,
        "margin_right": 2.0,
        "detailed_requirements": [
            "Sử dụng thông số mặc định của Bộ môn do chưa phân tích được qua AI."
        ],
    }

    if not pdf_text or not api_key:
        return default_rules

    prompt = f"""
Bạn là chuyên gia kiểm tra định dạng luận văn. Hãy đọc TOÀN BỘ văn bản quy định dưới đây và trích xuất thông số kỹ thuật chuẩn xác.

NỘI DUNG FILE QUY ĐỊNH PDF:
{pdf_text}

Hãy trả về DUY NHẤT một chuỗi JSON thuần có cấu trúc sau (không dùng mã markdown):
{{
  "font_name": "Times New Roman",
  "font_size": 13.0,
  "line_spacing": 1.5,
  "margin_top": 3.5,
  "margin_bottom": 3.0,
  "margin_left": 3.0,
  "margin_right": 2.0,
  "detailed_requirements": [
      "Trích xuất tất cả các quy định cụ thể tìm thấy trong PDF..."
  ]
}}
"""

    candidate_models = ["gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro"]

    for model_name in candidate_models:
        if genai is not None:
            try:
                client = genai.Client(api_key=api_key)
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt
                )
                if response and response.text:
                    raw_text = response.text.strip()
                    bt = chr(96) * 3
                    clean_json = raw_text.replace(f"{bt}json", "").replace(bt, "").strip()
                    return json.loads(clean_json)
            except Exception:
                pass

        if genai_legacy is not None:
            try:
                genai_legacy.configure(api_key=api_key)
                model = genai_legacy.GenerativeModel(model_name)
                response = model.generate_content(prompt)
                if response and response.text:
                    raw_text = response.text.strip()
                    bt = chr(96) * 3
                    clean_json = raw_text.replace(f"{bt}json", "").replace(bt, "").strip()
                    return json.loads(clean_json)
            except Exception:
                pass

    return default_rules

# ---------------------------------------------------------
# 4. HÀM TỐI ƯU & BẢO TOÀN TRANG BÌA (DÙNG ELEMENT XML GỐC)
# ---------------------------------------------------------
def optimize_cover_pages(doc):
    cover_errors = []
    cover_p_elements = set()

    cover_paragraphs = []

    # 1. Thu thập các đoạn văn trong Bảng khung bìa
    for table in doc.tables[:2]:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    cover_paragraphs.append(p)

    # 2. Thu thập các đoạn văn thuộc 2 trang đầu
    sections_p = [[]]
    for p in doc.paragraphs:
        sections_p[-1].append(p)
        xml = p._element.xml
        is_break = "\x0c" in p.text or ('w:br' in xml and 'type="page"' in xml) or ('w:sectPr' in xml)
        if is_break:
            sections_p.append([])
            if len(sections_p) > 2:
                break

    for idx_cover in range(min(2, len(sections_p))):
        for p in sections_p[idx_cover]:
            cover_paragraphs.append(p)

    # 3. Đánh dấu danh sách paragraph thuộc Trang Bìa
    for p in cover_paragraphs:
        cover_p_elements.add(p._element)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)

        if not p.text.strip():
            for run in p.runs:
                run.font.size = Pt(6)
            if not p.runs:
                r = p.add_run()
                r.font.size = Pt(6)

    cover_errors.append(
        "🖼️ **Bảo toàn Trang Bìa & Logo:** Đã giữ nguyên khung viền, logo trường, co nhỏ dòng trống (6pt) & cố định trang bìa gọn gàng!"
    )

    return cover_errors, cover_p_elements

# ---------------------------------------------------------
# 5. HÀM NHẬN BIẾT BỨC ẢNH CÓ PHẢI LÀ LOGO / TRANG BÌA KHÔNG
# ---------------------------------------------------------
def is_cover_or_logo_image(p_idx, p, cover_p_elements, paragraphs):
    if p._element in cover_p_elements:
        return True

    if p_idx < 30:
        context_text = ""
        start_k = max(0, p_idx - 6)
        end_k = min(len(paragraphs), p_idx + 7)
        for k in range(start_k, end_k):
            context_text += " " + paragraphs[k].text.upper()

        cover_keywords = [
            "BỘ GIÁO DỤC", "BỘ Y TẾ", "TRƯỜNG ĐẠI HỌC", "ĐẠI HỌC Y", 
            "LUẬN VĂN", "LUẬN ÁN", "KHÓA LUẬN", "BÁO CÁO", "TÊN ĐỀ TÀI",
            "NGƯỜI HƯỚNG DẪN", "HÀ NỘI"
        ]
        if any(kw in context_text for kw in cover_keywords):
            return True

    return False

# ---------------------------------------------------------
# 6. HÀM LẤY ĐỊNH DẠNG MẪU TỪ CÁC MỤC 2.1, 2.2, 2.3
# ---------------------------------------------------------
def get_reference_heading_format(doc):
    ref_pattern = re.compile(r"^\s*2\.[1-3]\.?", re.IGNORECASE)
    for p in doc.paragraphs:
        if ref_pattern.search(p.text.strip()):
            return {
                "left_indent": p.paragraph_format.left_indent,
                "first_line_indent": p.paragraph_format.first_line_indent,
                "alignment": p.paragraph_format.alignment,
                "space_before": p.paragraph_format.space_before,
                "space_after": p.paragraph_format.space_after,
            }
    return {
        "left_indent": Pt(0),
        "first_line_indent": Pt(0),
        "alignment": None,
        "space_before": Pt(3),
        "space_after": Pt(3),
    }

# ---------------------------------------------------------
# 7. HÀM SỬA & HIGHLIGHT MỤC 2.4 (VẤN ĐỀ ĐẠO ĐỨC)
# ---------------------------------------------------------
def fix_ethics_section(doc):
    detailed_errors = []
    ethics_pattern = re.compile(
        r"(2\.\d+\.?\s*)?(vấn đề đạo đức trong nghiên cứu|đạo đức trong nghiên cứu)", 
        re.IGNORECASE
    )
    ref_fmt = get_reference_heading_format(doc)

    def apply_heading_format(p):
        p.paragraph_format.left_indent = ref_fmt["left_indent"]
        p.paragraph_format.first_line_indent = ref_fmt["first_line_indent"]
        if ref_fmt["alignment"] is not None:
            p.paragraph_format.alignment = ref_fmt["alignment"]
        if ref_fmt["space_before"] is not None:
            p.paragraph_format.space_before = ref_fmt["space_before"]
        if ref_fmt["space_after"] is not None:
            p.paragraph_format.space_after = ref_fmt["space_after"]

    for idx, p in enumerate(list(doc.paragraphs)):
        text = p.text
        if not text.strip():
            continue

        match = ethics_pattern.search(text)
        if match:
            start_pos = match.start()
            matched_text = match.group(0)

            if start_pos > 0 and text[:start_pos].strip() != "":
                detailed_errors.append(
                    f"⚠️ **Đoạn {idx+1}:** Mục `{matched_text}` bị dính câu trước ➡️ **Đã ngắt dòng, in đậm, căn lề thẳng hàng 2.1-2.3 và HIGHLIGHT màu vàng.**"
                )
                text_before = text[:start_pos].rstrip()
                text_after = text[start_pos:].lstrip()

                p.text = text_before

                new_p = doc.add_paragraph()
                p._p.addnext(new_p._p)

                r_title = new_p.add_run(matched_text)
                r_title.bold = True
                r_title.font.highlight_color = WD_COLOR_INDEX.YELLOW
                apply_heading_format(new_p)

                rest_text = text_after[len(matched_text):]
                if rest_text:
                    r_rest = new_p.add_run(rest_text)
                    r_rest.bold = False

            else:
                apply_heading_format(p)
                detailed_errors.append(
                    f"📏 **Mục `{matched_text}`:** Đã căn thẳng hàng 100% với 2.1, 2.2, 2.3 và **HIGHLIGHT màu vàng**."
                )
                p.text = ""
                r_title = p.add_run(matched_text)
                r_title.bold = True
                r_title.font.highlight_color = WD_COLOR_INDEX.YELLOW
                rest_text = text[len(matched_text):]
                if rest_text:
                    r_rest = p.add_run(rest_text)
                    r_rest.bold = False

    return detailed_errors

# ---------------------------------------------------------
# 8. HÀM ÉP IN ĐẬM & HIGHLIGHT TÊN ĐỀ TÀI
# ---------------------------------------------------------
def enforce_bold_for_thesis_title(p):
    text_lower = p.text.lower().strip()
    is_title = (
        "tên đề tài" in text_lower
        or text_lower.startswith("đề tài:")
        or text_lower.startswith("đề tài :")
        or "tên đề tài:" in text_lower
    )
    if is_title:
        for run in p.runs:
            run.bold = True
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return True
    return False

# ---------------------------------------------------------
# ---------------------------------------------------------
# HÀM BẢO TOÀN ĐỊNH DẠNG KHI THAY ĐỔI VĂN BẢN TRONG ĐOẠN 
# ---------------------------------------------------------
def set_paragraph_text_preserve_formatting(p, new_text, highlight_changes=False, missing_year_ranges=None):
    char_map = []
    for run in p.runs:
        b = bool(run.bold) if run.bold is not None else False
        i = bool(run.italic) if run.italic is not None else False
        sup = bool(run.font.superscript) if run.font.superscript is not None else False
        sub = bool(run.font.subscript) if run.font.subscript is not None else False
        hl = run.font.highlight_color
        fn = run.font.name
        fs = run.font.size
        for ch in run.text:
            char_map.append((ch, b, i, sup, sub, fn, fs, hl))

    if not char_map or not new_text:
        p.text = new_text
        return

    matcher = difflib.SequenceMatcher(None, p.text, new_text)
    new_char_map = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            new_char_map.extend(char_map[i1:i2])
        elif tag == "replace":
            ref_fmt = list(char_map[i1] if i1 < len(char_map) else char_map[-1])
            if highlight_changes:
                ref_fmt[7] = WD_COLOR_INDEX.YELLOW  # Chỉ bôi vàng đoạn thay thế
            for _ in range(j2 - j1):
                new_char_map.append(tuple(ref_fmt))
        elif tag == "insert":
            ref_idx = max(0, i1 - 1)
            ref_fmt = list(char_map[ref_idx] if ref_idx < len(char_map) else char_map[-1])
            if highlight_changes:
                ref_fmt[7] = WD_COLOR_INDEX.YELLOW  # Chỉ bôi vàng đoạn thêm mới
            for _ in range(j2 - j1):
                new_char_map.append(tuple(ref_fmt))
        elif tag == "delete":
            pass

    # Bôi vàng chữ "và cộng sự" bị thiếu năm
    if missing_year_ranges:
        for start, end in missing_year_ranges:
            for idx in range(start, end):
                if idx < len(new_char_map):
                    fmt = list(new_char_map[idx])
                    fmt[7] = WD_COLOR_INDEX.YELLOW
                    new_char_map[idx] = tuple(fmt)

    p.text = ""
    if len(new_char_map) != len(new_text):
        p.text = new_text
        return

    curr_str = ""
    curr_fmt = new_char_map[0]

    for ch, fmt in zip(new_text, new_char_map):
        if fmt[1:] == curr_fmt[1:]:
            curr_str += ch
        else:
            r = p.add_run(curr_str)
            r.bold, r.italic, r.font.superscript, r.font.subscript = curr_fmt[1], curr_fmt[2], curr_fmt[3], curr_fmt[4]
            if curr_fmt[5]: r.font.name = curr_fmt[5]
            if curr_fmt[6]: r.font.size = curr_fmt[6]
            if curr_fmt[7]: r.font.highlight_color = curr_fmt[7]
            
            curr_str = ch
            curr_fmt = fmt

    if curr_str:
        r = p.add_run(curr_str)
        r.bold, r.italic, r.font.superscript, r.font.subscript = curr_fmt[1], curr_fmt[2], curr_fmt[3], curr_fmt[4]
        if curr_fmt[5]: r.font.name = curr_fmt[5]
        if curr_fmt[6]: r.font.size = curr_fmt[6]
        if curr_fmt[7]: r.font.highlight_color = curr_fmt[7]
# ---------------------------------------------------------
# 9. HÀM ĐỔI TRÍCH DẪN [1] THÀNH LŨY THỪA (BẢO TOÀN LŨY THỪA CÓ SẴN)
# ---------------------------------------------------------
def _add_mapped_runs(p, sub_text, sub_map, font_target, size_target):
    if not sub_text or not sub_map:
        return
    curr_str = ""
    curr_bold = sub_map[0][1]
    curr_italic = sub_map[0][2]
    curr_super = sub_map[0][3]
    curr_sub = sub_map[0][4]

    for ch, (_, b, i, sup, sub) in zip(sub_text, sub_map):
        if b == curr_bold and i == curr_italic and sup == curr_super and sub == curr_sub:
            curr_str += ch
        else:
            r = p.add_run(curr_str)
            r.bold = curr_bold
            r.italic = curr_italic
            r.font.superscript = curr_super
            r.font.subscript = curr_sub
            r.font.name = font_target
            if not curr_super:
                r.font.size = Pt(size_target)
            curr_str = ch
            curr_bold = b
            curr_italic = i
            curr_super = sup
            curr_sub = sub

    if curr_str:
        r = p.add_run(curr_str)
        r.bold = curr_bold
        r.italic = curr_italic
        r.font.superscript = curr_super
        r.font.subscript = curr_sub
        r.font.name = font_target
        if not curr_super:
            r.font.size = Pt(size_target)

def convert_brackets_to_superscript(p, font_target, size_target):
    text = p.text
    if not text or "[" not in text or "]" not in text:
        return False

    pattern = r"(\[\d+(?:[\s,–-]\d+)*\])"
    matches = list(re.finditer(pattern, text))
    if not matches:
        return False

    # Lấy bản đồ định dạng cho TẤT CẢ các thuộc tính bao gồm LŨY THỪA (superscript)
    char_map = []
    for run in p.runs:
        r_bold = bool(run.bold)
        r_italic = bool(run.italic)
        r_super = bool(run.font.superscript)
        r_sub = bool(run.font.subscript)
        for ch in run.text:
            char_map.append((ch, r_bold, r_italic, r_super, r_sub))

    if len(char_map) != len(text):
        char_map = [(ch, False, False, False, False) for ch in text]

    # Kiểm tra xem có trích dẫn ngoặc vuông nào CHƯA phải lũy thừa không
    needs_conversion = False
    for match in matches:
        start, end = match.span()
        if not all(char_map[k][3] for k in range(start, end)):
            needs_conversion = True
            break

    # Nếu tất cả trích dẫn ngoặc vuông ĐÃ LÀ LŨY THỪA RỒI -> Bỏ qua, bảo toàn nguyên vẹn
    if not needs_conversion:
        return False

    p.text = ""
    pos = 0
    for match in matches:
        start, end = match.span()
        if start > pos:
            _add_mapped_runs(p, text[pos:start], char_map[pos:start], font_target, size_target)

        num_str = match.group(1)[1:-1]
        match_bold = char_map[start][1] if start < len(char_map) else False
        match_italic = char_map[start][2] if start < len(char_map) else False

        r = p.add_run(num_str)
        r.font.superscript = True
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        r.bold = match_bold
        r.italic = match_italic
        r.font.name = font_target
        r.font.size = Pt(size_target)

        pos = end

    if pos < len(text):
        _add_mapped_runs(p, text[pos:], char_map[pos:], font_target, size_target)

    return True

# ---------------------------------------------------------
# 10. HÀM XÓA KHOẢNG TRẮNG DƯ THỪA
# ---------------------------------------------------------
def clean_spaces_and_punctuation(p):
    if not p.runs:
        return False

    text_upper = p.text.upper()
    cover_keywords = ["BỘ GIÁO DỤC", "BỘ Y TẾ", "TRƯỜNG ĐẠI HỌC", "VIỆN NGHIÊN CỨU", "UBND"]
    if any(keyword in text_upper for keyword in cover_keywords):
        return False

    changed = False
    for run in p.runs:
        if not run.text:
            continue
        orig = run.text
        cleaned = orig.replace("\xa0", " ").replace("\u200b", "")
        cleaned = re.sub(r" {2,}", " ", cleaned)
        cleaned = re.sub(r"\s+([:,.,!?,])", r"\1", cleaned)

        if cleaned != orig:
            run.text = cleaned
            changed = True

    runs_with_text = [r for r in p.runs if r.text]
    for idx in range(len(runs_with_text) - 1):
        curr_run = runs_with_text[idx]
        next_run = runs_with_text[idx + 1]
        while (
            curr_run.text
            and next_run.text
            and curr_run.text[-1] == " "
            and next_run.text[0] == " "
        ):
            next_run.text = next_run.text[1:]
            changed = True

    if runs_with_text:
        first_run = runs_with_text[0]
        last_run = runs_with_text[-1]
        l_stripped = first_run.text.lstrip(" \t\xa0")
        if l_stripped != first_run.text:
            first_run.text = l_stripped
            changed = True
        r_stripped = last_run.text.rstrip(" \t\xa0")
        if r_stripped != last_run.text:
            last_run.text = r_stripped
            changed = True

    return changed

# ---------------------------------------------------------
# 11. HÀM TỰ ĐỘNG BỔ SUNG "VÀ CỘNG SỰ" VÀ HIGHLIGHT
# ---------------------------------------------------------
def fix_author_citations(p, in_references_section=False):
    text = p.text
    if not text:
        return []

    changed = []
    original_text = text
    working_text = text

    # 1. Bỏ phần doi:... trong toàn bộ văn bản (kể cả Tài liệu tham khảo)
    if re.search(r"(?i)\bdoi:\s*\S+", working_text):
        working_text = re.sub(r"(?i)\s*\bdoi:\s*\S+", "", working_text).strip()
        changed.append("Đã xóa phần doi:")

    # 2. Xử lý "và cộng sự" (BỎ QUA phần TÀI LIỆU THAM KHẢO)
    if not in_references_section:
        if re.search(r"\bet\s+al\.?", working_text, re.IGNORECASE):
            working_text = re.sub(r"\bet\s+al\.?", "và cộng sự", working_text, flags=re.IGNORECASE)
            changed.append("Đã chuyển 'et al.' thành 'và cộng sự'")

        pattern_narrative = r"\b([A-ZÀ-Ỹ][a-zà-ỹA-Za-zÀ-Ỹ]+(?:\s+[A-ZÀ-Ỹ][a-zà-ỹA-Za-zÀ-Ỹ]+){0,3})[\s\xa0]*\(((?:19|20)\d{2})\)"
        pattern_parenthetical = r"\((?:bởi\s+)?([A-ZÀ-Ỹ][a-zà-ỹA-Za-zÀ-Ỹ]+(?:\s+[A-ZÀ-Ỹ][a-zà-ỹA-Za-zÀ-Ỹ]+){0,3})[\s\xa0]*,[\s\xa0]*((?:19|20)\d{2})\)"
        lead_words_pattern = r"^(Theo|Nghiên cứu của|Báo cáo của|Tác giả|Của|Trong)\s+"

        def repl_narrative(m):
            raw_name = m.group(1).strip()
            year = m.group(2)
            lead_match = re.match(lead_words_pattern, raw_name, re.IGNORECASE)
            lead_prefix = lead_match.group(0) if lead_match else ""
            clean_name = raw_name[len(lead_prefix):].strip() if lead_match else raw_name
            if not clean_name:
                clean_name, lead_prefix = raw_name, ""
            if re.search(r"(và\s+cộng\s+sự|và\s+cs\.?|et\s+al\.?)$", clean_name, re.IGNORECASE):
                return m.group(0)
            return f"{lead_prefix}{clean_name} và cộng sự ({year})"

        def repl_parenthetical(m):
            raw_name = m.group(1).strip()
            year = m.group(2)
            if re.search(r"(và\s+cộng\s+sự|và\s+cs\.?|et\s+al\.?)$", raw_name, re.IGNORECASE):
                return m.group(0)
            return f"({raw_name} và cộng sự, {year})"

        working_text = re.sub(pattern_narrative, repl_narrative, working_text)
        working_text = re.sub(pattern_parenthetical, repl_parenthetical, working_text)

    # 3. Quét kiểm tra "và cộng sự" bị thiếu năm
    missing_year_ranges = []
    if not in_references_section:
        for match in re.finditer(r"và cộng sự", working_text, re.IGNORECASE):
            start, end = match.span()
            lookahead = working_text[end:end+25]  # Nhìn phía trước 25 ký tự để tìm năm
            if not re.search(r"(?:19|20)\d{2}", lookahead):
                missing_year_ranges.append((start, end))

        if missing_year_ranges:
            changed.append("Đã bôi vàng 'và cộng sự' thiếu năm")

        if working_text != original_text and "và cộng sự" in working_text:
            changed.append("Đã sửa / thêm 'và cộng sự'")

    # Cập nhật lại văn bản nếu có sự thay đổi hoặc có vị trí thiếu năm
    if working_text != original_text or missing_year_ranges:
        set_paragraph_text_preserve_formatting(
            p, 
            working_text, 
            highlight_changes=True, 
            missing_year_ranges=missing_year_ranges
        )

    return changed

def has_image(paragraph):
    xml = paragraph._element.xml
    return "w:drawing" in xml or "w:pict" in xml or "a:blip" in xml

# ---------------------------------------------------------
# 12. ĐIỀU HÀNH CHUẨN HÓA TOÀN BỘ FILE DOCX
# ---------------------------------------------------------
def process_docx_file(uploaded_bytes, rules):
    doc = docx.Document(io.BytesIO(uploaded_bytes))

    cover_errors, cover_p_elements = optimize_cover_pages(doc)
    ethics_errors = fix_ethics_section(doc)

    detailed_errors = cover_errors + ethics_errors

    font_target = rules["font_name"]
    size_target = float(rules["font_size"])
    line_target = float(rules["line_spacing"])
    target_left = float(rules["margin_left"])
    target_right = float(rules["margin_right"])
    target_top = float(rules["margin_top"])
    target_bottom = float(rules["margin_bottom"])

    for idx, section in enumerate(doc.sections):
        curr_left = round(section.left_margin.cm, 1) if section.left_margin else 0
        curr_right = round(section.right_margin.cm, 1) if section.right_margin else 0
        curr_top = round(section.top_margin.cm, 1) if section.top_margin else 0
        curr_bottom = round(section.bottom_margin.cm, 1) if section.bottom_margin else 0
        if (
            curr_left != target_left
            or curr_right != target_right
            or curr_top != target_top
            or curr_bottom != target_bottom
        ):
            detailed_errors.append(
                f"📌 **Lề Trang (Phần {idx+1}):** Cũ ({curr_left}-{curr_right}-{curr_top}-{curr_bottom}cm) ➡️ **Đã sửa chuẩn quy định** ({target_left}-{target_right}-{target_top}-{target_bottom}cm)"
            )
        section.left_margin = Cm(target_left)
        section.right_margin = Cm(target_right)
        section.top_margin = Cm(target_top)
        section.bottom_margin = Cm(target_bottom)

    spacing_count = 0
    title_bold_count = 0
    fixed_authors = []
    converted_citation_count = 0
    paragraphs = doc.paragraphs

    for i, p in enumerate(paragraphs):
        if clean_spaces_and_punctuation(p):
            spacing_count += 1
        if enforce_bold_for_thesis_title(p):
            title_bold_count += 1
        authors = fix_author_citations(p)
        if authors:
            fixed_authors.extend(authors)
        if convert_brackets_to_superscript(p, font_target, size_target):
            converted_citation_count += 1

        # KIỂM TRA HÌNH ẢNH (ĐÃ BỔ SUNG QUÉT PHẠM VI RỘNG & TỪ KHÓA NGOÙN DẠNG WEBSITE/LINK/TRÍCH DẪN)
        if has_image(p):
            if is_cover_or_logo_image(i, p, cover_p_elements, paragraphs):
                continue

            # Mở rộng quét 3 đoạn trước và 4 đoạn sau hình ảnh
            start_idx = max(0, i - 3)
            end_idx = min(len(paragraphs), i + 5)
            nearby_text = " ".join([paragraphs[k].text.strip() for k in range(start_idx, end_idx) if paragraphs[k].text.strip()]).lower()

            missing_text = []

            # 1. Kiểm tra chú thích hình ảnh
            caption_keywords = ["hình", "sơ đồ", "biểu đồ", "đồ thị", "bản đồ", "ảnh", "figure", "fig.", "chart"]
            if not any(kw in nearby_text for kw in caption_keywords):
                missing_text.append("Thiếu Tên/Chú thích hình")

            # 2. Kiểm tra nguồn trích dẫn / URL / Website
            source_keywords = [
                "nguồn", "source", "website", "web", "link", "trích", "theo", "tác giả", 
                "tự chụp", "tổng hợp", "bộ y tế", "who", "cdc", "bệnh viện", "bv", 
                "http", "https", "www.", ".com", ".vn", ".org", ".edu", ".gov"
            ]
            has_numeric_citation = bool(re.search(r"\[\d+\]|\(\d+\)|(?:19|20)\d{2}", nearby_text))
            has_source_keyword = any(kw in nearby_text for kw in source_keywords)

            if not (has_source_keyword or has_numeric_citation):
                missing_text.append("Thiếu Nguồn trích dẫn")

            if missing_text:
                err_msg = " & ".join(missing_text)
                detailed_errors.append(f"🖼️ **Hình ảnh tại đoạn {i+1}:** {err_msg}")

                warn_p = doc.add_paragraph()
                p._p.addnext(warn_p._p)
                r_warn = warn_p.add_run(f"⚠️ [CẢNH BÁO ĐỊNH DẠNG]: Bức ảnh này đang {err_msg}!")
                r_warn.font.highlight_color = WD_COLOR_INDEX.YELLOW
                r_warn.bold = True

    if title_bold_count > 0:
        detailed_errors.append("🖋️ **Tên Đề Tài:** Đã ép in đậm và **HIGHLIGHT màu vàng**.")
    if spacing_count > 0:
        detailed_errors.append(f"🧹 **Dấu Cách Thừa:** Đã dọn dẹp tại **{spacing_count} đoạn** văn.")
    if fixed_authors:
        detailed_errors.append(f"✍️ **Trích Dẫn Tác Giả:** Đã bổ sung 'và cộng sự' và **HIGHLIGHT màu vàng** tại **{len(fixed_authors)} vị trí**.")
    if converted_citation_count > 0:
        detailed_errors.append(f"✨ **Trích Dẫn Số:** Đã chuyển **{converted_citation_count} đoạn** ngoặc `[1]` thành lũy thừa `¹` và **HIGHLIGHT màu vàng**.")

    # Áp dụng định dạng chuẩn cho Thân bài (Bỏ qua Trang Bìa và giữ nguyên lũy thừa)
    for p in paragraphs:
        if not p.text.strip():
            continue
        if p._element not in cover_p_elements:
            p.paragraph_format.line_spacing = line_target
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = font_target
                if not run.font.superscript:
                    run.font.size = Pt(size_target)

    # Áp dụng định dạng cho Bảng (Bỏ qua Bảng Bìa)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p._element in cover_p_elements:
                        continue
                    clean_spaces_and_punctuation(p)
                    convert_brackets_to_superscript(p, font_target, size_target)
                    p.paragraph_format.line_spacing = line_target
                    for run in p.runs:
                        run.font.name = font_target
                        if not run.font.superscript:
                            run.font.size = Pt(size_target)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output, detailed_errors

# ---------------------------------------------------------
# 13. GIAO DIỆN CHÍNH STREAMLIT
# ---------------------------------------------------------
col_main, col_side = st.columns([2, 1])

with col_main:
    st.subheader("1. Chọn File Quy Định Trình Bày (.PDF)")
    pdf_folder = "quy_dinh"
    pdf_files = (
        [f for f in os.listdir(pdf_folder) if f.endswith(".pdf")]
        if os.path.exists(pdf_folder)
        else []
    )
    selected_pdf = (
        st.selectbox("Chọn file quy định PDF trong thư mục `quy_dinh`:", pdf_files)
        if pdf_files
        else None
    )

    pdf_text = ""
    parsed_rules = {
        "font_name": "Times New Roman",
        "font_size": 13.0,
        "line_spacing": 1.5,
        "margin_top": 3.5,
        "margin_bottom": 3.0,
        "margin_left": 3.0,
        "margin_right": 2.0,
        "detailed_requirements": [],
    }

    if selected_pdf:
        pdf_path = os.path.join(pdf_folder, selected_pdf)
        pdf_text = extract_raw_text_from_pdf(pdf_path)
        if pdf_text:
            parsed_rules = analyze_rules_with_gemini(pdf_text, API_KEY)
            st.success(f"✅ Đã đọc thành công 100% nội dung file `{selected_pdf}`!")

            with st.expander("📌 Báo cáo các quy định AI đã trích xuất từ file PDF"):
                st.markdown(f"**Font chữ:** {parsed_rules.get('font_name')} | **Cỡ chữ:** {parsed_rules.get('font_size')}pt | **Giãn dòng:** {parsed_rules.get('line_spacing')}")
                st.markdown(f"**Lề (Trái-Phải-Trên-Dưới):** {parsed_rules.get('margin_left')}cm - {parsed_rules.get('margin_right')}cm - {parsed_rules.get('margin_top')}cm - {parsed_rules.get('margin_bottom')}cm")
                st.markdown("**Các yêu cầu chi tiết đọc được từ PDF:**")
                reqs = parsed_rules.get("detailed_requirements", [])
                if isinstance(reqs, list):
                    for req in reqs:
                        st.markdown(f"* {req}")
                else:
                    st.write(reqs)
        else:
            st.warning("⚠️ File PDF dạng Scan/Ảnh chụp. Bạn có thể tùy chỉnh thông số ở Sidebar bên phải!")

    st.subheader("2. Tải File Luận Văn Của Học Viên (.DOCX)")
    uploaded_docx = st.file_uploader("Thả file .docx luận văn vào đây", type=["docx"])

# ---------------------------------------------------------
# 14. THANH BÊN (SIDEBAR)
# ---------------------------------------------------------
st.sidebar.title("📄 Tải Template Mẫu")
template_path = os.path.join("template", "template.docx")
if os.path.exists(template_path):
    with open(template_path, "rb") as f:
        template_bytes = f.read()
    st.sidebar.download_button(
        label="📥 TẢI TEMPLATE WORD MẪU (.DOCX)",
        data=template_bytes,
        file_name="template.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
    st.sidebar.caption("👉 Tải file mẫu về để điền nội dung chuẩn định dạng.")
else:
    st.sidebar.info("ℹ️ Không tìm thấy file `template.docx` trong thư mục `template/`.")

st.sidebar.markdown("---")
st.sidebar.title("⚙️ Bảng Quy Định Định Dạng")

final_font = st.sidebar.text_input("Font Chữ Thân Bài", value=parsed_rules.get("font_name", "Times New Roman"))
final_size = st.sidebar.number_input("Cỡ Chữ (pt)", value=float(parsed_rules.get("font_size", 13.0)), step=0.5)
final_spacing = st.sidebar.number_input("Giãn Dòng (Line Spacing)", value=float(parsed_rules.get("line_spacing", 1.5)), step=0.1)

st.sidebar.markdown("---")
st.sidebar.subheader("📐 Căn Lề Trang (cm)")
final_left = st.sidebar.number_input("Lề Trái (Left)", value=float(parsed_rules.get("margin_left", 3.0)), step=0.5)
final_right = st.sidebar.number_input("Lề Phải (Right)", value=float(parsed_rules.get("margin_right", 2.0)), step=0.5)
final_top = st.sidebar.number_input("Lề Trên (Top)", value=float(parsed_rules.get("margin_top", 3.5)), step=0.5)
final_bottom = st.sidebar.number_input("Lề Dưới (Bottom)", value=float(parsed_rules.get("margin_bottom", 3.0)), step=0.5)

active_rules = {
    "font_name": final_font,
    "font_size": final_size,
    "line_spacing": final_spacing,
    "margin_left": final_left,
    "margin_right": final_right,
    "margin_top": final_top,
    "margin_bottom": final_bottom,
}

# ---------------------------------------------------------
# 15. NÚT BẤM CHẠY CHƯƠNG TRÌNH
# ---------------------------------------------------------
with col_main:
    st.markdown("---")
    if st.button("🔍 KIỂM TRA ĐỊNH DẠNG", type="primary", use_container_width=True):
        if not uploaded_docx:
            st.error("❌ Vui lòng tải file Luận Văn (.docx) ở Bước 2 trước!")
        else:
            with st.spinner("⏳ Đang bảo toàn logo/bìa, sửa lỗi định dạng & HIGHLIGHT màu vàng..."):
                fixed_stream, error_list = process_docx_file(
                    uploaded_docx.getvalue(), active_rules
                )
            st.markdown("### 📋 BẢNG BÁO CÁO KẾT QUẢ KIỂM TRA & SỬA LỖI")
            if error_list:
                for err in error_list:
                    st.write(err)
                st.success("🎉 **Đã hoàn thành! Toàn bộ vị trí sửa đổi đã được HIGHLIGHT màu vàng trong file Word.**")
            else:
                st.success("🎉 File luận văn đã hoàn toàn đạt chuẩn định dạng!")

            st.download_button(
                label="📥 TẢI FILE LUẬN VĂN ĐÃ ĐƯỢC CHUẨN HÓA & HIGHLIGHT",
                data=fixed_stream,
                file_name="LuanVan_DaiHocYHaNoi_DaChuanHoa.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
